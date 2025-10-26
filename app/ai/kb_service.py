"""Knowledge Base service for document processing and RAG."""
import os
import time
from pathlib import Path
from datetime import datetime
from typing import List, Optional, Dict, Any, Tuple
from uuid import UUID

from sqlalchemy import select, func, and_, or_
from sqlalchemy.orm import Session
from langchain_openai import OpenAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import (
    TextLoader,
    PyPDFLoader,
    Docx2txtLoader,
)

from app.config import settings
from app.models.knowledge_base import KnowledgeBaseDocument, KnowledgeBaseChunk
from app.schemas.knowledge_base import (
    KBSearchResultChunk,
    KBSearchResponse,
    KBStatsResponse,
)


class KBService:
    """Service for Knowledge Base operations."""
    
    def __init__(self, db: Session):
        """Initialize KB service."""
        self.db = db
        self.embeddings = OpenAIEmbeddings(
            api_key=settings.OPENAI_API_KEY,
            model=settings.EMBEDDING_MODEL
        )
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )
    
    async def process_document(
        self,
        document_id: UUID,
        file_path: str
    ) -> bool:
        """
        Process document: extract text, chunk, and generate embeddings.
        
        Args:
            document_id: Document UUID
            file_path: Path to uploaded file
            
        Returns:
            Success status
        """
        print(f"🔄 Processing document: {document_id}")
        print(f"📁 File path: {file_path}")
        
        try:
            # Get document from DB
            document = self.db.get(KnowledgeBaseDocument, document_id)
            if not document:
                print(f"❌ Document not found: {document_id}")
                return False
            
            print(f"✅ Document found: {document.original_filename}")
            
            # Update status to processing
            document.status = "processing"
            self.db.commit()
            print("📝 Status updated to 'processing'")
            
            # Load and extract text based on file type
            print(f"📄 Extracting text from {document.content_type}...")
            text_content = await self._extract_text(file_path, document.content_type)
            
            if not text_content:
                print("❌ No text content extracted!")
                document.status = "failed"
                self.db.commit()
                return False
            
            print(f"✅ Extracted {len(text_content)} characters")
            
            # Split into chunks
            chunks = self.text_splitter.split_text(text_content)
            print(f"✂️  Split into {len(chunks)} chunks")
            
            # Generate embeddings and store chunks
            for idx, chunk_text in enumerate(chunks):
                print(f"🧮 Generating embedding for chunk {idx + 1}/{len(chunks)}...")
                
                # Generate embedding
                embedding = await self.embeddings.aembed_query(chunk_text)
                print(f"✅ Embedding generated: {len(embedding)} dimensions")
                
                # Create chunk
                chunk = KnowledgeBaseChunk(
                    document_id=document_id,
                    chunk_index=idx,
                    content=chunk_text,
                    char_count=len(chunk_text),
                    token_count=self._estimate_tokens(chunk_text),
                    embedding=embedding
                )
                self.db.add(chunk)
            
            # Update document
            document.total_chunks = len(chunks)
            document.status = "completed"
            document.processed_at = datetime.utcnow()
            
            self.db.commit()
            print(f"✅ Document processing complete! {len(chunks)} chunks saved")
            return True
            
        except Exception as e:
            print(f"❌ Error processing document {document_id}: {e}")
            import traceback
            traceback.print_exc()
            
            if document:
                document.status = "failed"
                self.db.commit()
            return False
    
    async def _extract_text(self, file_path: str, content_type: str) -> Optional[str]:
        """Extract text from file based on type."""
        try:
            path = Path(file_path)
            
            print(f"📂 Checking file: {path}")
            print(f"📝 Content type: {content_type}")
            
            if not path.exists():
                print(f"❌ File not found: {path}")
                return None
            
            print(f"✅ File exists, size: {path.stat().st_size} bytes")
            
            # PDF files
            if content_type == "application/pdf" or path.suffix.lower() == ".pdf":
                print("📕 Processing as PDF...")
                loader = PyPDFLoader(str(path))
                documents = loader.load()
                text = "\n\n".join([doc.page_content for doc in documents])
                print(f"✅ PDF extracted: {len(text)} chars, {len(documents)} pages")
                return text
            
            # Word documents
            elif content_type in [
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "application/msword"
            ] or path.suffix.lower() in [".docx", ".doc"]:
                print("📘 Processing as Word document...")
                try:
                    from docx import Document as DocxDocument
                    
                    # Try with python-docx directly for better compatibility
                    doc = DocxDocument(str(path))
                    
                    # Extract text from all paragraphs
                    full_text = []
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            full_text.append(paragraph.text)
                    
                    # Also extract from tables
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = []
                            for cell in row.cells:
                                if cell.text.strip():
                                    row_text.append(cell.text)
                            if row_text:
                                full_text.append(" | ".join(row_text))
                    
                    text = "\n\n".join(full_text)
                    print(f"✅ Word doc extracted: {len(text)} chars, {len(doc.paragraphs)} paragraphs")
                    return text
                    
                except Exception as docx_error:
                    print(f"⚠️  python-docx failed: {docx_error}, trying Docx2txtLoader...")
                    try:
                        loader = Docx2txtLoader(str(path))
                        documents = loader.load()
                        text = "\n\n".join([doc.page_content for doc in documents])
                        print(f"✅ Word doc extracted (fallback): {len(text)} chars")
                        return text
                    except Exception as loader_error:
                        print(f"❌ Both DOCX methods failed: {loader_error}")
                        raise
            
            # Text files
            elif content_type.startswith("text/") or path.suffix.lower() in [
                ".txt", ".md", ".csv", ".json", ".xml", ".html", ".css", ".js", ".py"
            ]:
                print("📄 Processing as text file...")
                loader = TextLoader(str(path), encoding="utf-8")
                documents = loader.load()
                text = "\n\n".join([doc.page_content for doc in documents])
                print(f"✅ Text file extracted: {len(text)} chars")
                return text
            
            else:
                # Try as text file
                print("❓ Unknown type, trying as text...")
                with open(path, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read()
                print(f"✅ Generic text extracted: {len(text)} chars")
                return text
                    
        except Exception as e:
            print(f"❌ Error extracting text from {file_path}: {e}")
            import traceback
            traceback.print_exc()
            return None
    
    def _estimate_tokens(self, text: str) -> int:
        """Estimate token count (rough approximation)."""
        # Rough estimate: ~4 characters per token
        return len(text) // 4
    
    async def search(
        self,
        query: str,
        k: int = 5,
        category: Optional[str] = None,
        user_id: Optional[UUID] = None
    ) -> KBSearchResponse:
        """
        Semantic search in knowledge base.
        
        Args:
            query: Search query
            k: Number of results
            category: Filter by category
            user_id: Filter by user
            
        Returns:
            Search results
        """
        start_time = time.time()
        
        try:
            # Generate query embedding
            query_embedding = await self.embeddings.aembed_query(query)
            
            # Build query with filters
            # Use pgvector's cosine distance operator <=>
            from sqlalchemy import literal_column
            
            stmt = select(
                KnowledgeBaseChunk,
                KnowledgeBaseDocument,
                # Calculate cosine distance using pgvector operator <=>
                (KnowledgeBaseChunk.embedding.l2_distance(query_embedding)).label("distance")
            ).join(
                KnowledgeBaseDocument,
                KnowledgeBaseChunk.document_id == KnowledgeBaseDocument.id
            ).where(
                KnowledgeBaseDocument.status == "completed"
            )
            
            # Apply filters
            if category:
                stmt = stmt.where(KnowledgeBaseDocument.category == category)
            if user_id:
                stmt = stmt.where(KnowledgeBaseDocument.user_id == user_id)
            
            # Order by similarity and limit
            stmt = stmt.order_by("distance").limit(k)
            
            # Execute query
            results = self.db.execute(stmt).all()
            
            # Format results
            search_results = []
            for chunk, document, distance in results:
                # Convert distance to similarity score (0-1)
                similarity_score = 1 - distance if distance else 0
                
                search_results.append(
                    KBSearchResultChunk(
                        chunk_id=chunk.id,
                        document_id=document.id,
                        content=chunk.content,
                        similarity_score=similarity_score,
                        chunk_index=chunk.chunk_index,
                        filename=document.original_filename,
                        category=document.category,
                        created_at=document.created_at
                    )
                )
            
            processing_time = (time.time() - start_time) * 1000  # Convert to ms
            
            return KBSearchResponse(
                query=query,
                results=search_results,
                total_results=len(search_results),
                processing_time_ms=round(processing_time, 2)
            )
            
        except Exception as e:
            print(f"Error searching KB: {e}")
            return KBSearchResponse(
                query=query,
                results=[],
                total_results=0,
                processing_time_ms=0
            )
    
    async def get_context_for_query(
        self,
        query: str,
        k: int = 3,
        user_id: Optional[UUID] = None
    ) -> str:
        """
        Get relevant context for a query (for RAG).
        
        Args:
            query: User query
            k: Number of chunks to retrieve
            user_id: Filter by user
            
        Returns:
            Formatted context string
        """
        search_result = await self.search(query=query, k=k, user_id=user_id)
        
        if not search_result.results:
            return ""
        
        # Format context
        context_parts = []
        for idx, result in enumerate(search_result.results, 1):
            context_parts.append(
                f"[Source {idx}: {result.filename}]\n{result.content}\n"
            )
        
        return "\n".join(context_parts)
    
    def get_document(self, document_id: UUID) -> Optional[KnowledgeBaseDocument]:
        """Get document by ID."""
        return self.db.get(KnowledgeBaseDocument, document_id)
    
    def list_documents(
        self,
        user_id: Optional[UUID] = None,
        category: Optional[str] = None,
        status: Optional[str] = None,
        skip: int = 0,
        limit: int = 100
    ) -> Tuple[List[KnowledgeBaseDocument], int]:
        """
        List documents with filters.
        
        Returns:
            Tuple of (documents, total_count)
        """
        # Build query
        stmt = select(KnowledgeBaseDocument)
        
        # Apply filters
        filters = []
        if user_id:
            filters.append(KnowledgeBaseDocument.user_id == user_id)
        if category:
            filters.append(KnowledgeBaseDocument.category == category)
        if status:
            filters.append(KnowledgeBaseDocument.status == status)
        
        if filters:
            stmt = stmt.where(and_(*filters))
        
        # Get total count
        count_stmt = select(func.count()).select_from(stmt.subquery())
        total = self.db.execute(count_stmt).scalar() or 0
        
        # Apply pagination and order
        stmt = stmt.order_by(KnowledgeBaseDocument.created_at.desc())
        stmt = stmt.offset(skip).limit(limit)
        
        # Execute
        documents = self.db.execute(stmt).scalars().all()
        
        return list(documents), total
    
    def delete_document(self, document_id: UUID) -> bool:
        """Delete document and all its chunks."""
        try:
            document = self.db.get(KnowledgeBaseDocument, document_id)
            if not document:
                return False
            
            # Delete file from filesystem
            try:
                file_path = Path(document.file_path)
                if file_path.exists():
                    file_path.unlink()
            except Exception as e:
                print(f"Error deleting file: {e}")
            
            # Delete from database (chunks will cascade)
            self.db.delete(document)
            self.db.commit()
            
            return True
        except Exception as e:
            print(f"Error deleting document: {e}")
            self.db.rollback()
            return False
    
    def get_stats(self, user_id: Optional[UUID] = None) -> KBStatsResponse:
        """Get KB statistics."""
        # Build base query
        doc_stmt = select(KnowledgeBaseDocument)
        if user_id:
            doc_stmt = doc_stmt.where(KnowledgeBaseDocument.user_id == user_id)
        
        # Total documents
        total_docs = self.db.execute(
            select(func.count()).select_from(doc_stmt.subquery())
        ).scalar() or 0
        
        # Total chunks
        chunk_stmt = select(func.count(KnowledgeBaseChunk.id))
        if user_id:
            chunk_stmt = chunk_stmt.join(
                KnowledgeBaseDocument
            ).where(KnowledgeBaseDocument.user_id == user_id)
        total_chunks = self.db.execute(chunk_stmt).scalar() or 0
        
        # Total size
        size_stmt = select(func.sum(KnowledgeBaseDocument.file_size))
        if user_id:
            size_stmt = size_stmt.where(KnowledgeBaseDocument.user_id == user_id)
        total_size = self.db.execute(size_stmt).scalar() or 0
        
        # Documents by category
        cat_stmt = select(
            KnowledgeBaseDocument.category,
            func.count(KnowledgeBaseDocument.id)
        )
        if user_id:
            cat_stmt = cat_stmt.where(KnowledgeBaseDocument.user_id == user_id)
        cat_stmt = cat_stmt.group_by(KnowledgeBaseDocument.category)
        
        docs_by_category = dict(self.db.execute(cat_stmt).all())
        
        # Documents by status
        status_stmt = select(
            KnowledgeBaseDocument.status,
            func.count(KnowledgeBaseDocument.id)
        )
        if user_id:
            status_stmt = status_stmt.where(KnowledgeBaseDocument.user_id == user_id)
        status_stmt = status_stmt.group_by(KnowledgeBaseDocument.status)
        
        docs_by_status = dict(self.db.execute(status_stmt).all())
        
        # Recent uploads
        recent_stmt = doc_stmt.order_by(
            KnowledgeBaseDocument.created_at.desc()
        ).limit(5)
        recent_docs = list(self.db.execute(recent_stmt).scalars().all())
        
        return KBStatsResponse(
            total_documents=total_docs,
            total_chunks=total_chunks,
            total_size_bytes=int(total_size),
            documents_by_category=docs_by_category,
            documents_by_status=docs_by_status,
            recent_uploads=recent_docs
        )

